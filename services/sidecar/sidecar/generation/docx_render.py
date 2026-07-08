import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Mm, Pt
from docxtpl import DocxTemplate, InlineImage

from sidecar.generation.context_keys import caption_key, image_key, section_key
from sidecar.models.report_input import ReportInput
from sidecar.models.template import Template

# Opening/closing quote chars built with chr() to avoid smart-quote encoding issues in source file
_Q_OPEN = {chr(0x201C), chr(0x201E), chr(0xAB), chr(0x201A), chr(0x22)}   # “ „ « ‚ “
_Q_CLOSE = {chr(0x201D), chr(0xBB), chr(0x2019), chr(0x22)}  # “ » ‘ “
_ELLIPSIS_BRACKET = chr(91) + chr(46) * 3 + chr(93)  # “[...]”


_MOJIBAKE = {
    chr(0xE2) + chr(0x20AC) + chr(0x0153): chr(0x201C),   # " LEFT DOUBLE QUOTATION MARK
    chr(0xE2) + chr(0x20AC) + chr(0x009D): chr(0x201D),   # " RIGHT DOUBLE QUOTATION MARK
    chr(0xE2) + chr(0x20AC) + chr(0x2122): chr(0x2019),   # ' RIGHT SINGLE QUOTATION MARK
    chr(0xE2) + chr(0x20AC) + chr(0x02DC): chr(0x2018),   # ' LEFT SINGLE QUOTATION MARK
    chr(0xE2) + chr(0x20AC) + chr(0x201C): chr(0x2013),   # – EN DASH
    chr(0xE2) + chr(0x20AC) + chr(0x201D): chr(0x2014),   # — EM DASH
    chr(0xC2) + chr(0xBB): chr(0xBB),                     # » RIGHT-POINTING DOUBLE ANGLE QUOTATION
}


def _sanitize_text(text: str) -> str:
    """Fix text extracted from PDF before writing to DOCX XML:
    - CP1252 mojibake: PDF bytes decoded byte-by-byte via CP1252 instead of UTF-8
    - Lone surrogates (U+D800..U+DFFF) from corrupt PDF encoding
    - Soft hyphens (U+00AD) that appear as literal chars in DOCX text
    """
    for seq, replacement in _MOJIBAKE.items():
        text = text.replace(seq, replacement)
    cleaned = []
    for ch in text:
        cp = ord(ch)
        if 0xD800 <= cp <= 0xDFFF:
            if cp == 0xDC9D:
                cleaned.append(chr(0x201D))  # RIGHT DOUBLE QUOTATION MARK
            # else drop lone surrogate
        elif cp == 0x00AD:
            pass  # strip soft hyphens
        else:
            cleaned.append(ch)
    return "".join(cleaned)


def _resolve_inline_vars(text: str, var_context: dict[str, str]) -> str:
    """Substitute {{key}} placeholders inside section text with actual variable values."""
    for key, value in var_context.items():
        text = text.replace(f"{{{{{key}}}}}", value)
    return text


def _inject_var_placeholders(text: str, variables) -> str:
    """Replace detected source values in text with {{key}} tokens (mirrors JS injectPlaceholders)."""
    for var in variables:
        val = var.source_value_detected or ""
        if len(val) < 6:
            continue
        text = re.sub(re.escape(val), f"{{{{{var.key}}}}}", text)
    return text


def build_render_context(
    template: Template,
    report_input: ReportInput,
    section_overrides: dict[str, str] | None = None,
) -> dict[str, str]:
    """Builds the Jinja2 context dict shared by DOCX rendering and HTML
    preview, so both always reflect the same underlying data."""

    context: dict[str, str] = {}

    var_context: dict[str, str] = {}
    for var in template.variables:
        value = next((v.value for v in report_input.variables if v.variable_id == var.id), "")
        context[var.key] = value
        var_context[var.key] = value

    for section in template.sections:
        text = next((s.text for s in report_input.sections if s.section_id == section.id), "")
        if section_overrides and section.id in section_overrides:
            text = section_overrides[section.id]
        # Strip figure-caption lines that were extracted from the PDF body text.
        # The skeleton already renders each caption via {{ caption_key }} below the image,
        # so keeping them here produces a duplicate paragraph.
        text = "\n".join(
            line for line in text.split("\n")
            if not re.match(r"^\s*Figura\s+\d+", line, re.IGNORECASE)
        )
        resolved = _resolve_inline_vars(_sanitize_text(text), var_context)
        # \a = docxtpl paragraph break; \n in stored text = paragraph separator
        context[section_key(section)] = resolved.replace("\n", "\a")

    for placeholder in template.image_placeholders:
        label = _inject_var_placeholders(_sanitize_text(placeholder.label), template.variables)
        context[caption_key(placeholder)] = _resolve_inline_vars(label, var_context)

    return context


def render_docx(
    template: Template,
    report_input: ReportInput,
    output_path: Path,
    section_overrides: dict[str, str] | None = None,
) -> Path:
    if not template.docx_skeleton_path:
        raise ValueError("Template has no docx_skeleton_path; confirm the template first.")

    tpl = DocxTemplate(template.docx_skeleton_path)
    context = build_render_context(template, report_input, section_overrides)

    for placeholder in template.image_placeholders:
        images = [img for img in report_input.images if img.placeholder_id == placeholder.id]
        if images:
            context[image_key(placeholder)] = InlineImage(tpl, images[0].file_path, width=Mm(120))
        else:
            context[image_key(placeholder)] = ""

    var_context: dict[str, str] = {
        var.key: next((v.value for v in report_input.variables if v.variable_id == var.id), "")
        for var in template.variables
    }

    tpl.render(context)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    tpl.save(output_path)
    _postprocess_captions(output_path, template, var_context)
    _postprocess_paragraphs(output_path)
    _postprocess_signature(output_path)
    return output_path


def _is_quote_start(text: str) -> bool:
    return bool(text) and text[0] in _Q_OPEN


def _is_quote_end(text: str) -> bool:
    if not text:
        return False
    stripped = text.rstrip()
    # Last non-space char is a closing quote
    if stripped and stripped[-1] in _Q_CLOSE:
        return True
    # Ends with "[...]" (Brazilian legal truncation marker) optionally followed by 1-2 chars
    return stripped.endswith(_ELLIPSIS_BRACKET) or _ELLIPSIS_BRACKET in stripped[-8:]


_HEADING_RE = re.compile(r'^\d+\.\s')
_CAPTION_RE = re.compile(r'^Figura\s+\d+', re.IGNORECASE)
_LIST_ITEM_RE = re.compile(r'^[·•–—•·\-]\s|^Vest[ií]gio\s+\S+', re.UNICODE)
_SIGNATURE_RE = re.compile(
    r'^(.+?)\s+Perito\s+Criminal\s+Mat\.\s*([\d-]+)\s*$',
    re.IGNORECASE | re.UNICODE,
)


def _postprocess_captions(docx_path: Path, template, var_context: dict[str, str]) -> None:
    """Fix figure captions: substitute variable values AND ensure caption is BELOW the image.

    Handles two skeleton generations:
    - New skeleton (image first, then caption): text substitution only.
    - Old skeleton (caption first, then image): swaps order + substitutes text.
    """
    from docx.oxml.ns import qn

    doc = Document(str(docx_path))
    changed = False

    # Pass 1: text substitution in caption paragraphs
    for p in doc.paragraphs:
        text = p.text.strip()
        if not _CAPTION_RE.match(text):
            continue
        new_text = _inject_var_placeholders(text, template.variables)
        new_text = _resolve_inline_vars(new_text, var_context)
        if new_text != text and p.runs:
            p.runs[0].text = new_text
            for run in p.runs[1:]:
                run.text = ""
            changed = True

    # Pass 2: fix order — caption before image → move caption after image
    body_el = doc.element.body
    children = list(body_el)
    # Collect (caption_el, image_el) pairs where caption comes immediately before an image paragraph
    pairs = []
    for i in range(len(children) - 1):
        el = children[i]
        if el.tag != qn("w:p"):
            continue
        text = "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()
        if not _CAPTION_RE.match(text):
            continue
        next_el = children[i + 1]
        if next_el.tag != qn("w:p"):
            continue
        if next_el.find(f".//{qn('w:drawing')}") is not None:
            pairs.append((el, next_el))

    for cap_el, img_el in pairs:
        body_el.remove(cap_el)
        new_children = list(body_el)
        img_idx = new_children.index(img_el)
        body_el.insert(img_idx + 1, cap_el)
        changed = True

    if changed:
        doc.save(str(docx_path))


def _postprocess_signature(docx_path: Path) -> None:
    """Detect 'NAME Perito Criminal Mat. XXXX' paragraphs and reformat as 3 centered lines.

    Result:
        [space 24pt]
        NAME                  ← centered, Arial 12pt
        Perito Criminal       ← centered, Arial 12pt
        Mat. XXXX             ← centered, Arial 11pt
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document(str(docx_path))
    body = doc.element.body
    body_children = list(body)
    changed = False

    for i, el in enumerate(body_children):
        if el.tag != qn("w:p"):
            continue
        text = "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()
        m = _SIGNATURE_RE.match(text)
        if not m:
            continue

        name = m.group(1).strip()
        mat = m.group(2).strip()

        def _sig_para(line: str, size_pt: int, space_before_pt: int = 0) -> OxmlElement:
            pe = OxmlElement("w:p")
            pPr = OxmlElement("w:pPr")
            jc = OxmlElement("w:jc")
            jc.set(qn("w:val"), "center")
            pPr.append(jc)
            ind = OxmlElement("w:ind")
            ind.set(qn("w:firstLine"), "0")
            pPr.append(ind)
            sp = OxmlElement("w:spacing")
            sp.set(qn("w:before"), str(space_before_pt * 20))
            sp.set(qn("w:after"), "0")
            sp.set(qn("w:line"), "360")
            sp.set(qn("w:lineRule"), "auto")
            pPr.append(sp)
            pe.append(pPr)
            r = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")
            rf = OxmlElement("w:rFonts")
            rf.set(qn("w:ascii"), "Arial")
            rf.set(qn("w:hAnsi"), "Arial")
            rPr.append(rf)
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), str(size_pt * 2))
            rPr.append(sz)
            r.append(rPr)
            t = OxmlElement("w:t")
            t.text = line
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            r.append(t)
            pe.append(r)
            return pe

        body.remove(el)
        # Insert in reverse order so they appear in correct reading order
        body.insert(i, _sig_para(f"Mat. {mat}", 11))
        body.insert(i, _sig_para("Perito Criminal", 12))
        body.insert(i, _sig_para(name, 12, space_before_pt=24))
        changed = True
        break  # Only one signature block per document

    if changed:
        doc.save(str(docx_path))


def _postprocess_paragraphs(docx_path: Path) -> None:
    """Apply formatting that docxtpl cannot control per-paragraph.

    - All body paragraphs: JUSTIFY + 1.5 line spacing (docxtpl paragraph cloning
      sometimes drops the jc/spacing XML from the skeleton pPr).
    - Block quotes: left_indent=5cm, italic, no first-line indent.
    - List items (bullet · or "Vestígio N"): no first-line indent, left flush.
    """
    doc = Document(str(docx_path))
    in_quote = False

    for p in doc.paragraphs:
        text = p.text.strip()

        if not in_quote and _is_quote_start(text):
            in_quote = True

        if in_quote:
            p.paragraph_format.left_indent = Cm(5.0)
            p.paragraph_format.right_indent = Pt(0)
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.italic = True
                run.font.name = "Arial"
                run.font.size = Pt(12)

            if _is_quote_end(text):
                in_quote = False

        elif _HEADING_RE.match(text) or not text:
            # Section headings: enforce compact spacing + explicit Arial
            if _HEADING_RE.match(text):
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                p.paragraph_format.line_spacing = 1.5
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(12)

        elif _CAPTION_RE.match(text):
            # Figure captions: center, no indent, compact — leave alignment/bold as set
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.first_line_indent = Pt(0)
            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = Pt(11)

        else:
            # Body paragraph: enforce justify + 1.5 spacing + explicit Arial on every run
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf = p.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = 1.5
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            # List items (bullet or "Vestígio N:") should NOT carry first-line indent
            if _LIST_ITEM_RE.match(text) or text.startswith(chr(0x00B7)):
                pf.first_line_indent = Pt(0)
            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = Pt(12)

    doc.save(str(docx_path))
