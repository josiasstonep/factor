from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Mm, Pt

from sidecar.generation.context_keys import caption_key, image_key, section_key
from sidecar.models.template import Template, TemplateImagePlaceholder

_CONTENT_WIDTH = Mm(210) - Cm(2.0) - Cm(2.0)  # A4 minus left+right margins


def _apply_global_style(doc: Document) -> None:
    # Normal: Arial 12pt, 1.5 spacing, justified, first-line indent 1.25cm
    try:
        normal = doc.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(12)
        pf = normal.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.5
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.first_line_indent = Cm(1.25)
    except KeyError:
        pass
    # "No Spacing" style: must stay single-spaced (used by captions, headers, etc.)
    for ns_name in ("No Spacing", "Sem Espaçamento", "Sem Espacamento"):
        try:
            ns = doc.styles[ns_name]
            ns.font.name = "Arial"
            ns.font.size = Pt(12)
            pf = ns.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            break
        except KeyError:
            pass
    # Heading styles: compact spacing, no first-line indent
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        try:
            style = doc.styles[style_name]
            style.font.name = "Arial"
            style.font.size = Pt(12)
            pf = style.paragraph_format
            pf.first_line_indent = Pt(0)
            pf.space_before = Pt(6)
            pf.space_after = Pt(0)
        except KeyError:
            pass


def _setup_page(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_height = Mm(297)
    sec.page_width = Mm(210)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    sec.top_margin = Cm(2.8)       # body starts 2.8cm below top (below header image)
    sec.bottom_margin = Cm(2.0)
    sec.header_distance = Cm(0.43) # header image starts 0.43cm from top edge
    sec.footer_distance = Pt(0)    # no wasted footer margin


def _add_header_image(doc: Document, image_path: str) -> None:
    sec = doc.sections[0]
    header = sec.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Extend into left margin so the image spans the full A4 page width
    p.paragraph_format.left_indent = -sec.left_margin
    p.paragraph_format.right_indent = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(image_path, width=sec.page_width)  # full A4 width (210mm)


def _add_page_numbers(doc: Document) -> None:
    """Add centered 'Página X de Y' to the footer, on a new paragraph after the image (if any)."""
    sec = doc.sections[0]
    footer = sec.footer

    # Add a new paragraph for page numbers (after footer image paragraph)
    p = doc.add_paragraph() if False else footer.add_paragraph()  # footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)

    def _add_field(run_elem, fld_char_type: str | None = None, instr: str | None = None):
        if fld_char_type:
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), fld_char_type)
            run_elem.append(fc)
        if instr:
            it = OxmlElement("w:instrText")
            it.set(qn("xml:space"), "preserve")
            it.text = instr
            run_elem.append(it)

    run = p.add_run("Página ")
    run.font.name = "Arial"
    run.font.size = Pt(9)

    r_page = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    fn = OxmlElement("w:rFonts")
    fn.set(qn("w:ascii"), "Arial")
    fn.set(qn("w:hAnsi"), "Arial")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")  # 9pt = 18 half-points
    rpr.append(fn)
    rpr.append(sz)
    r_page.append(rpr)
    _add_field(r_page, "begin")
    _add_field(r_page, instr=" PAGE ")
    _add_field(r_page, "end")
    p.runs[-1]._r.addnext(r_page)

    run2 = p.add_run(" de ")
    run2.font.name = "Arial"
    run2.font.size = Pt(9)

    r_pages = OxmlElement("w:r")
    rpr2 = OxmlElement("w:rPr")
    fn2 = OxmlElement("w:rFonts")
    fn2.set(qn("w:ascii"), "Arial")
    fn2.set(qn("w:hAnsi"), "Arial")
    sz2 = OxmlElement("w:sz")
    sz2.set(qn("w:val"), "18")
    rpr2.append(fn2)
    rpr2.append(sz2)
    r_pages.append(rpr2)
    _add_field(r_pages, "begin")
    _add_field(r_pages, instr=" NUMPAGES ")
    _add_field(r_pages, "end")
    run2._r.addnext(r_pages)


def _add_footer_image(doc: Document, image_path: str) -> None:
    sec = doc.sections[0]
    footer = sec.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = -sec.left_margin
    p.paragraph_format.right_indent = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(image_path, width=sec.page_width)  # full A4 width (210mm)


def _heading_para(doc: Document, text: str, order: int = -1) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent = Pt(0)
    label = f"{order + 1}.  {text.upper()}" if order >= 0 else text.upper()
    run = p.add_run(label)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(12)


def _body_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Cm(1.25)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(12)


def _add_image_placeholder(doc: Document, placeholder: TemplateImagePlaceholder) -> None:
    # Image first, caption below (standard figure format)
    img_para = doc.add_paragraph("{{ " + image_key(placeholder) + " }}")
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.paragraph_format.first_line_indent = Pt(0)
    img_para.paragraph_format.space_before = Pt(12)
    img_para.paragraph_format.space_after = Pt(4)
    caption = doc.add_paragraph()
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(12)
    caption.paragraph_format.first_line_indent = Pt(0)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = caption.add_run("{{ " + caption_key(placeholder) + " }}")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(11)


def build_skeleton(template: Template, output_path: Path) -> Path:
    doc = Document()
    _apply_global_style(doc)
    _setup_page(doc)

    if template.header_image_path:
        _add_header_image(doc, template.header_image_path)

    if template.footer_image_path:
        _add_footer_image(doc, template.footer_image_path)

    _add_page_numbers(doc)

    # Group image placeholders by section_id; images without section_id go at the end
    images_by_section: dict[str, list[TemplateImagePlaceholder]] = defaultdict(list)
    orphan_images: list[TemplateImagePlaceholder] = []
    for p in sorted(template.image_placeholders, key=lambda p: p.order):
        if p.section_id:
            images_by_section[p.section_id].append(p)
        else:
            orphan_images.append(p)

    for section in sorted(template.sections, key=lambda s: s.order):
        _heading_para(doc, section.label, section.order)
        _body_para(doc, "{{ " + section_key(section) + " }}")
        for placeholder in images_by_section.get(section.id, []):
            _add_image_placeholder(doc, placeholder)

    # Orphan placeholders (no section_id) appended at the end
    for placeholder in orphan_images:
        _add_image_placeholder(doc, placeholder)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path
